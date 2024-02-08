from bs4 import BeautifulSoup
import pyperclip
import re
from docx import Document


def clean_html_tags(raw_html):
    clean_text = BeautifulSoup(raw_html, "html.parser").text
    return clean_text


doc = Document()

# Load HTML from the Proof file exported from Twine
with open(
    "C:\\Users\\Sarah\\Downloads\\Advanced Creative Writing Sample.html",
    "r",
    encoding="utf8",
) as file:
    html_content = file.read()

# Parse the HTML
soup = BeautifulSoup(html_content, "html.parser")

# Find all <tw-passagedata> elements
passages = soup.find_all("tw-passagedata")

text_to_copy = ""
content_for_word_count = ""
word_count = 0

# Extract and accumulate the name and contents of each passage
for passage in passages:
    content = passage.text.strip()

    if content is None or content == "":
        continue

    title = passage.get("name")

    res = len(re.findall(r"\w+", content))
    word_count += int(res)

    text_to_copy += (
        title + "\n" + content + "\n"
    )  # Add passage name and content to the accumulated text
    content_for_word_count += content

    doc.add_heading(title, level=1)
    doc.add_paragraph(content)

# Copy all the accumulated text to the clipboard
pyperclip.copy(text_to_copy)

to_remove = [
    "Advanced Creative Writing Sample",
    '<<back "←">>',
    '<<linkreplace "',
    '" t8n>>',
    "<</linkreplace>>",
]

for t in to_remove:
    content = content.replace(t, " ")

content_for_word_count = re.sub(r"\"[a-z_]*\">><</link>>", " ", content_for_word_count)

content_for_word_count = re.sub(r"\[", " ", content_for_word_count)
content_for_word_count = re.sub(r"\]", " ", content_for_word_count)

content_for_word_count = clean_html_tags(content_for_word_count)

content_for_word_count = re.sub(r"[ ]+", " ", content_for_word_count)
content_for_word_count = re.sub(r">", "", content_for_word_count)
content_for_word_count = re.sub(r"<", "", content_for_word_count)
content_for_word_count = re.sub(r"\\", " ", content_for_word_count)

# For testing
# print(word_count_content)
# pyperclip.copy(word_count_content)

word_count = len(content_for_word_count.strip().split())

doc.save("C:\\Users\\Sarah\\Downloads\\output.docx")

print(f"Word Count: {word_count} words")
